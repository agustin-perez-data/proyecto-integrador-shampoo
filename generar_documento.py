"""
Genera el documento Word del Proyecto Integrador.
Ejecutar desde la raiz del proyecto: python generar_documento.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUTPUT = Path(__file__).parent / "data" / "Proyecto_Integrador_Shampoo.docx"

# ── Colores ──────────────────────────────────────────────────────────────────
AZUL_OSC  = RGBColor(0x1A, 0x1A, 0x2E)
AZUL_MED  = RGBColor(0x1A, 0x52, 0x76)
AZUL_CLAR = RGBColor(0xEA, 0xF2, 0xFF)
VERDE     = RGBColor(0x1E, 0x84, 0x49)
AMARILLO  = RGBColor(0xFF, 0xEE, 0xBA)
GRIS_TEXT = RGBColor(0x44, 0x44, 0x44)
BLANCO    = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_color: str):
    """Pinta el fondo de una celda con color hexadecimal (sin #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_table_border(table):
    """Borde fino en todas las celdas de la tabla."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "C5D8EF")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_heading(doc, text, level=1, color=None, space_before=12, space_after=4):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = color or AZUL_OSC
    run.font.bold = True
    return p


def add_para(doc, text, italic=False, color=None, space_before=0, space_after=6, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_hipotesis(doc, tag, text):
    """Bloque azul claro con etiqueta + texto de hipótesis."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, "EAF2FF")
    cell.width = Cm(16)

    # Tag
    tp = cell.paragraphs[0]
    tp.paragraph_format.space_before = Pt(4)
    tp.paragraph_format.space_after  = Pt(2)
    tr = tp.add_run(tag)
    tr.font.bold  = True
    tr.font.size  = Pt(9)
    tr.font.color.rgb = AZUL_MED

    # Texto
    bp = cell.add_paragraph(text)
    bp.paragraph_format.space_before = Pt(0)
    bp.paragraph_format.space_after  = Pt(4)
    bp.runs[0].font.size = Pt(10.5)

    doc.add_paragraph()


def add_conclusion_block(doc, verdict, verdict_color, bg_hex, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_hex)

    vp = cell.paragraphs[0]
    vp.paragraph_format.space_before = Pt(4)
    vp.paragraph_format.space_after  = Pt(2)
    vr = vp.add_run(verdict)
    vr.font.bold  = True
    vr.font.size  = Pt(9)
    vr.font.color.rgb = verdict_color

    bp = cell.add_paragraph(text)
    bp.paragraph_format.space_before = Pt(0)
    bp.paragraph_format.space_after  = Pt(4)
    bp.runs[0].font.size = Pt(10.5)

    doc.add_paragraph()


def add_rec(doc, icon, text_bold, text_rest):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.add_run(f"{icon}  ").font.size = Pt(11)
    rb = p.add_run(text_bold)
    rb.font.bold = True
    rb.font.size = Pt(10.5)
    p.add_run(text_rest).font.size = Pt(10.5)


def add_annex_item(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, "F4F4F4")
    set_table_border(table)

    tp = cell.paragraphs[0]
    tp.paragraph_format.space_before = Pt(3)
    tp.paragraph_format.space_after  = Pt(1)
    tr = tp.add_run(title)
    tr.font.bold  = True
    tr.font.size  = Pt(10)
    tr.font.color.rgb = AZUL_MED

    bp = cell.add_paragraph(body)
    bp.paragraph_format.space_before = Pt(0)
    bp.paragraph_format.space_after  = Pt(4)
    bp.runs[0].font.size = Pt(9.5)
    bp.runs[0].font.color.rgb = GRIS_TEXT

    doc.add_paragraph()


# ═════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()

    # ── Márgenes ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # Estilo base del documento
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ══════════════════════════════════════════════════════════════════
    # PORTADA
    # ══════════════════════════════════════════════════════════════════
    inst = doc.add_paragraph("Universidad Argentina de la Empresa — Ciencia de Datos")
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst.runs[0].font.size  = Pt(10)
    inst.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    inst.paragraph_format.space_after = Pt(10)

    titulo = doc.add_heading("Análisis de Precios de Shampoos y Acondicionadores en el Mercado Digital Argentino", 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].font.size  = Pt(18)
    titulo.runs[0].font.color.rgb = AZUL_OSC
    titulo.paragraph_format.space_after = Pt(6)

    sub = doc.add_paragraph("Proyecto Integrador — 2026")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    sub.paragraph_format.space_after = Pt(14)

    for line in ["Alumno: Agustín Pérez",
                 "Dataset: 1.562 productos únicos disponibles · Jumbo, Farmacity, Disco",
                 "Período de scraping: Junio 2026"]:
        pm = doc.add_paragraph(line)
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pm.runs[0].font.size = Pt(10)
        pm.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        pm.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    # 1. HIPÓTESIS
    # ══════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Hipótesis de Análisis", level=1, color=AZUL_OSC)

    add_para(doc,
        "El mercado de cuidado capilar en Argentina presenta una oferta amplia y segmentada entre "
        "retailers con distintos perfiles de precio. Se plantean tres hipótesis sobre la estructura "
        "de precios en el canal digital:")

    add_hipotesis(doc,
        "H1 — Línea de producto",
        "Los productos de línea profesional tienen un precio por ml superior a los de línea estándar, "
        "debido a su formulación diferenciada y posicionamiento premium.")

    add_hipotesis(doc,
        "H2 — Tipo de cabello",
        "Los productos para cabello rizado o tratado/teñido presentan un mayor precio por ml que los "
        "de uso general, dado que están orientados a necesidades específicas con ingredientes especializados.")

    add_hipotesis(doc,
        "H3 — Canal de venta",
        "La dispersión de precios en Farmacity (farmacia especializada) es significativamente mayor "
        "que en Jumbo y Disco (supermercados), dado su mix de productos masivos y premium.")

    # ══════════════════════════════════════════════════════════════════
    # 2. METODOLOGÍA
    # ══════════════════════════════════════════════════════════════════
    add_heading(doc, "2. Descripción Metodológica", level=1, color=AZUL_OSC)

    # 2.1 Scraping
    add_heading(doc, "2.1  Scraping de datos", level=2, color=AZUL_MED, space_before=6)

    add_para(doc,
        "Se utilizó la API pública VTEX (plataforma de e-commerce en la que operan los tres retailers) "
        "para extraer productos de las categorías shampoo y acondicionador. El endpoint consultado fue:")

    code = doc.add_paragraph(
        "https://{retailer}/api/catalog_system/pub/products/search/{término}?_from=0&_to=49")
    code.runs[0].font.name = "Courier New"
    code.runs[0].font.size = Pt(9)
    code.paragraph_format.left_indent  = Cm(1)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after  = Pt(6)

    add_para(doc,
        "Se aplicó paginación automática (parámetros _from / _to) hasta recorrer el catálogo completo. "
        "Solo se retuvieron registros con disponible = True para garantizar que los precios reflejen "
        "la oferta activa al momento del scraping. Los términos de búsqueda incluyeron denominaciones "
        "genéricas (shampoo, acondicionador, tratamiento capilar) y nombres de marcas principales.")

    # Tabla de fuentes
    tbl = doc.add_table(rows=5, cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Fuente", "Tipo", "Registros raw", "Disponibles", "Dataset final"]
    data = [
        ["Jumbo",      "Supermercado", "2.631", "393",   "386"],
        ["Farmacity",  "Farmacia",     "1.111", "894",   "876"],
        ["Disco",      "Supermercado", "2.738", "304",   "300"],
        ["Total",      "",             "6.480", "1.591", "1.562"],
    ]
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        set_cell_bg(cell, "1A5276")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold  = True
        run.font.size  = Pt(10)
        run.font.color.rgb = BLANCO

    for i, row in enumerate(data):
        bg = "F4F8FF" if i % 2 == 0 else "FFFFFF"
        is_total = (i == 3)
        for j, val in enumerate(row):
            cell = tbl.cell(i + 1, j)
            set_cell_bg(cell, "EAF2FF" if is_total else bg)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.bold = is_total

    set_table_border(tbl)
    doc.add_paragraph()

    # 2.2 Limpieza
    add_heading(doc, "2.2  Limpieza y procesamiento", level=2, color=AZUL_MED, space_before=6)

    add_para(doc,
        "El pipeline de limpieza se implementó en Python (pandas) y aplicó los siguientes filtros "
        "secuencialmente:")

    pasos = [
        "Filtro de disponibilidad: eliminación de registros con disponible = False (precios históricos del catálogo VTEX sin stock activo).",
        "Filtro de precio mínimo: descarte de productos con precio inferior a $500 ARS.",
        "Eliminación de no capilares: remoción de productos colados por las búsquedas (juguetes, cremas corporales, colonias bebé) mediante lista de palabras clave negativas.",
        "Corrección de volumen: valores de volumen_ml superiores a 5.000 ml reasignados a nulo (errores de extracción regex del nombre).",
        "Deduplicación: eliminación de registros duplicados por combinación nombre + fuente.",
        "Normalización de marcas: unificación de case (PANTENE → Pantene) y variantes tipográficas.",
    ]
    for paso in pasos:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.runs[0].text if p.runs else p.add_run(paso)
        if not p.runs:
            p.add_run(paso).font.size = Pt(10.5)
        else:
            p.runs[0].font.size = Pt(10.5)

    add_para(doc,
        "Las variables linea_tipo (profesional / estándar) y tipo_cabello (general / rizado / tratado / "
        "anticaspa / seco_dañado / bebé / graso) se derivaron mediante clasificación por palabras clave "
        "en el nombre del producto.",
        space_before=6)

    # 2.3 Campos / DER
    add_heading(doc, "2.3  Campos de la base — DER y Diccionario de Datos", level=2, color=AZUL_MED, space_before=6)

    add_para(doc,
        "El dataset corresponde a una entidad única (producto disponible en retailer), con clave "
        "primaria compuesta por sku_id + fuente:")

    dic_rows = [
        ("sku_id",        "Entero",  "PK",  "ID único del SKU en el sistema VTEX del retailer"),
        ("producto_id",   "Entero",  "FK",  "ID del producto padre (agrupa variantes)"),
        ("fuente",        "Texto",   "—",   "Retailer de origen: Jumbo, Farmacity o Disco"),
        ("nombre",        "Texto",   "—",   "Nombre completo del producto (tal como aparece en el sitio)"),
        ("marca",         "Texto",   "—",   "Marca normalizada del producto"),
        ("precio_ars",    "Decimal", "—",   "Precio de venta en pesos argentinos (ARS)"),
        ("precio_lista",  "Decimal", "—",   "Precio de lista original antes de descuentos"),
        ("disponible",    "Bool",    "—",   "TRUE = disponible para compra al momento del scraping"),
        ("volumen_ml",    "Decimal", "—",   "Volumen en ml extraído del nombre con regex. NULL si no se encontró"),
        ("precio_por_ml", "Decimal", "—",   "precio_ars / volumen_ml. NULL si no hay volumen"),
        ("categoria_raw", "Texto",   "—",   "Categoría según el árbol del retailer"),
        ("tipo_producto", "Texto",   "—",   "shampoo / acondicionador / tratamiento (clasificado por Python)"),
        ("linea_tipo",    "Texto",   "—",   "profesional / estándar (clasificado por Python)"),
        ("tipo_cabello",  "Texto",   "—",   "general / rizado / tratado / anticaspa / seco_dañado / bebé / graso"),
        ("url",           "Texto",   "—",   "URL del producto en el sitio del retailer"),
    ]

    tbl2 = doc.add_table(rows=len(dic_rows) + 1, cols=4)
    tbl2.style = "Table Grid"
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["Campo", "Tipo", "Clave", "Descripción"]):
        cell = tbl2.cell(0, j)
        set_cell_bg(cell, "1A5276")
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = BLANCO

    widths = [Cm(3.2), Cm(1.8), Cm(1.2), Cm(9.5)]
    for i, (campo, tipo, clave, desc) in enumerate(dic_rows):
        bg = "F4F8FF" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate([campo, tipo, clave, desc]):
            cell = tbl2.cell(i + 1, j)
            cell.width = widths[j]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if j == 0:
                run.font.bold = True
                run.font.name = "Courier New"
            if j == 2 and val != "—":
                run.font.bold  = True
                run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    set_table_border(tbl2)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════
    # 3. CONCLUSIONES
    # ══════════════════════════════════════════════════════════════════
    add_heading(doc, "3. Conclusiones y Recomendaciones de Marketing", level=1, color=AZUL_OSC)

    add_heading(doc, "H1 — Línea profesional vs. estándar", level=2, color=AZUL_MED, space_before=4)
    add_conclusion_block(doc,
        verdict="H1 No confirmada",
        verdict_color=RGBColor(0xC8, 0x7F, 0x0A),
        bg_hex="FFF3CD",
        text="La mediana del precio por ml es prácticamente igual entre ambas líneas: $28,69/ml (profesional) "
             "vs. $28,31/ml (estándar). La media estándar es mayor ($84,88 vs. $43,40) por la presencia de "
             "outliers en esa categoría. En el mercado digital argentino, la etiqueta 'profesional' no implica "
             "necesariamente un precio por ml superior.")

    add_heading(doc, "H2 — Precio por ml según tipo de cabello", level=2, color=AZUL_MED, space_before=4)
    add_conclusion_block(doc,
        verdict="H2 Confirmada parcialmente",
        verdict_color=RGBColor(0x11, 0x7A, 0x65),
        bg_hex="D1ECF1",
        text="Los productos anticaspa presentan el mayor precio por ml promedio ($116,33/ml), seguidos por "
             "uso general ($78,08/ml). Contrariamente a la hipótesis, los productos para cabello rizado "
             "($27,33/ml) y graso ($22,77/ml) son los más económicos. La especialización no se traduce "
             "uniformemente en mayor precio: los anticaspa (menor volumen, formulación activa) presentan "
             "el mayor precio relativo.")

    add_heading(doc, "H3 — Dispersión de precios por fuente", level=2, color=AZUL_MED, space_before=4)
    add_conclusion_block(doc,
        verdict="H3 Confirmada",
        verdict_color=RGBColor(0x1E, 0x84, 0x49),
        bg_hex="D4EDDA",
        text="Farmacity tiene un coeficiente de variación del 115,7% (precio_ars), casi el triple que "
             "Jumbo (43,0%) y Disco (41,3%). Su catálogo mezcla productos masivos (desde $900 ARS) con "
             "kits premium (hasta $195.532 ARS). Los supermercados presentan una oferta más homogénea. "
             "Para consumidores con presupuesto definido, Jumbo o Disco ofrecen mayor previsibilidad.")

    add_heading(doc, "Regresión lineal — volumen_ml → precio_ars", level=2, color=AZUL_MED, space_before=4)
    add_conclusion_block(doc,
        verdict="Correlación débil — modelo explicativo limitado",
        verdict_color=RGBColor(0x11, 0x7A, 0x65),
        bg_hex="D1ECF1",
        text="La correlación entre volumen y precio es r = 0,117 (débil positiva), con un R² = 0,014. "
             "El volumen del envase explica apenas el 1,4% de la variación del precio. La ecuación "
             "resultante es: precio = $9.731 + $4,72 × volumen_ml. El precio se determina principalmente "
             "por la marca, la línea y el retailer, no por el volumen del envase.")

    add_heading(doc, "Recomendaciones de marketing", level=2, color=AZUL_MED, space_before=4)

    add_rec(doc, "🛒", "Para consumidores: ",
            "Comparar precio por ml antes de comprar. Jumbo y Disco ofrecen los mismos productos "
            "a precios promedio un 37–60% menores que Farmacity. La línea estándar tiene precio/ml "
            "equivalente a la profesional.")

    add_rec(doc, "🏪", "Para retailers (Jumbo / Disco): ",
            "Ampliar el surtido de productos anticaspa y de tratamiento capilar — categorías con mayor "
            "precio/ml donde Farmacity domina actualmente. Esto permitiría capturar consumidores "
            "dispuestos a pagar más por soluciones específicas.")

    add_rec(doc, "📊", "Para marcas: ",
            "El diferencial de precio entre línea profesional y estándar no está siendo percibido en "
            "el canal digital. Revisar la estrategia de pricing online para que el premium de la línea "
            "profesional sea visible al consumidor.")

    add_rec(doc, "🔍", "Limitaciones del análisis: ",
            "La clasificación de linea_tipo y tipo_cabello fue inferida por palabras clave en el nombre "
            "del producto, sin datos estructurados del retailer, lo que puede introducir errores. La "
            "regresión lineal simple sobre volumen tiene bajo poder explicativo; futuros análisis "
            "deberían incorporar marca, retailer y tipo_cabello como variables predictoras.")

    # ══════════════════════════════════════════════════════════════════
    # ANEXO
    # ══════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "Anexo — Material complementario", level=1, color=RGBColor(0x55, 0x55, 0x55))

    add_para(doc,
        "El siguiente material se encuentra en la planilla planilla_shampoo.xlsx, disponible en el "
        "repositorio GitHub del proyecto. Este anexo no cuenta dentro de las 2-3 carillas del cuerpo.",
        italic=True, color=GRIS_TEXT, size=10)

    add_annex_item(doc,
        "Tablas dinámicas (5) — hoja estadistica_descriptiva",
        "T1: Métricas globales de precio_ars (19 indicadores)  ·  T2: Métricas globales de precio_por_ml  ·  "
        "T3: precio_por_ml por linea_tipo — H1  ·  T4: precio_por_ml por tipo_cabello — H2  ·  "
        "T5: precio_ars por fuente (Jumbo / Farmacity / Disco) — H3")

    add_annex_item(doc,
        "Gráficos (8) — hoja visualizaciones",
        "G1: Distribución por fuente (torta)  ·  G2: Precio promedio por fuente (barras, H3)  ·  "
        "G3: Precio/ml profesional vs estándar (columnas, H1)  ·  G4: Precio/ml por tipo de cabello (barras, H2)  ·  "
        "G5: Histograma de precio_ars  ·  G6: Top 10 marcas por cantidad de productos (barras)  ·  "
        "G7: Tipo de producto por fuente (columnas apiladas)  ·  G8: Dispersión precio vs volumen")

    add_annex_item(doc,
        "Regresión lineal (1) — hoja regresion_lineal",
        "Variables: volumen_ml (X) → precio_ars (Y)  ·  n = 1.401 observaciones  ·  "
        "Fórmulas Sheets: CORREL, SLOPE, INTERCEPT, RSQ, FORECAST  ·  "
        "Scatter chart con línea de tendencia + ecuación + R²  ·  "
        "Tabla de predicción para 9 volúmenes (100–1500 ml)")

    # Pie del documento
    doc.add_paragraph()
    pie = doc.add_paragraph(
        "Proyecto Integrador — Ciencia de Datos · UADE 2026 · Agustín Pérez · "
        "Datos: Jumbo, Farmacity, Disco vía API VTEX · Junio 2026")
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.runs[0].font.size = Pt(8.5)
    pie.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Documento generado: {OUTPUT}")


if __name__ == "__main__":
    build()
